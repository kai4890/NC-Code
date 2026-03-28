"""
Document loader for AutoComply.

Handles PDF and DOCX file ingestion, raw text extraction, clause splitting,
and text normalisation.  Returns a clean list of clause strings ready for
semantic embedding.
"""

import os
import re
import tempfile
import logging
from pathlib import Path
from typing import List

import fitz          # PyMuPDF
from docx import Document

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def load_document(file_path: str) -> List[str]:
    """Load a PDF or DOCX file from disk and return clean clause strings.

    Args:
        file_path: Absolute or relative path to the PDF or DOCX file.

    Returns:
        List of clean, normalised clause strings (minimum 5 words each).

    Raises:
        FileNotFoundError: If the file does not exist.
        ValueError: If the file extension is not .pdf or .docx.
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    suffix = path.suffix.lower()
    if suffix == ".pdf":
        raw_text = _extract_pdf_text(str(path))
    elif suffix == ".docx":
        raw_text = _extract_docx_text(str(path))
    else:
        raise ValueError(
            f"Unsupported file format '{suffix}'. Only PDF and DOCX are accepted."
        )

    clauses = _split_into_clauses(raw_text)
    clean = [_clean_clause(c) for c in clauses]
    # Keep only clauses with enough content to be meaningful
    return [c for c in clean if len(c.split()) >= 5]


def load_document_from_bytes(file_bytes: bytes, filename: str) -> List[str]:
    """Load a document from raw bytes (Streamlit file-uploader compatible).

    Writes bytes to a temporary file, processes it with :func:`load_document`,
    then removes the temp file.

    Args:
        file_bytes: Raw bytes of the uploaded file.
        filename:   Original filename — used to determine the file type.

    Returns:
        List of clean, normalised clause strings.

    Raises:
        ValueError: If the file extension is not .pdf or .docx.
    """
    suffix = Path(filename).suffix.lower()
    if suffix not in (".pdf", ".docx"):
        raise ValueError(
            f"Unsupported file format '{suffix}'. Only PDF and DOCX are accepted."
        )

    tmp_fd, tmp_path = tempfile.mkstemp(suffix=suffix)
    try:
        with os.fdopen(tmp_fd, "wb") as tmp_file:
            tmp_file.write(file_bytes)
        clauses = load_document(tmp_path)
    finally:
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)

    return clauses


# ---------------------------------------------------------------------------
# Private extraction helpers
# ---------------------------------------------------------------------------

def _extract_pdf_text(file_path: str) -> str:
    """Extract raw text from every page of a PDF using PyMuPDF."""
    parts: List[str] = []
    with fitz.open(file_path) as doc:
        for page_num, page in enumerate(doc, start=1):
            text = page.get_text("text")
            if text.strip():
                parts.append(text)
            logger.debug(f"PDF page {page_num}: extracted {len(text)} chars")
    return "\n".join(parts)


def _extract_docx_text(file_path: str) -> str:
    """Extract raw text from a DOCX file, including paragraphs and tables."""
    doc = Document(file_path)
    parts: List[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # Tables often contain policy clauses in Word-based security documents
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    parts.append(text)

    return "\n".join(parts)


# ---------------------------------------------------------------------------
# Private splitting helpers
# ---------------------------------------------------------------------------

def _split_into_clauses(text: str) -> List[str]:
    """Split raw extracted text into individual clause strings.

    Strategy:
    1. Normalise line endings.
    2. Split on sentence-ending punctuation (. ! ?) followed by whitespace.
    3. Split on double newlines (paragraph breaks).
    4. Split on semicolons — common in policy documents.
    """
    # Normalise line endings
    text = text.replace("\r\n", "\n").replace("\r", "\n")

    # Split on sentence boundaries or double blank lines
    raw_sentences = re.split(r"(?<=[.!?])\s+|\n{2,}", text)

    # Further split on semicolons (policy clauses often use them)
    clauses: List[str] = []
    for sentence in raw_sentences:
        sub = sentence.split(";")
        clauses.extend(sub)

    return clauses


def _clean_clause(clause: str) -> str:
    """Normalise and clean a single extracted clause string.

    Removes:
    - Page number patterns  (e.g. "Page 1 of 20", "- 3 -")
    - Version / date stamps (e.g. "v2.1", "DRAFT")
    - Classification markings (UNCLASSIFIED, PROTECTED, etc.)
    - Excessive whitespace
    - Standalone numbers

    Normalises:
    - Unicode dashes → ASCII hyphen
    - Unicode smart quotes → ASCII quotes
    """
    # Page number patterns
    clause = re.sub(r"\bPage\s+\d+\s+of\s+\d+\b", "", clause, flags=re.IGNORECASE)
    clause = re.sub(r"^\s*[-–—]\s*\d+\s*[-–—]\s*$", "", clause, flags=re.MULTILINE)
    clause = re.sub(r"\b\d+\s*/\s*\d+\b", "", clause)

    # Version strings
    clause = re.sub(r"\bv\d+\.\d+\b", "", clause, flags=re.IGNORECASE)

    # Common document classification / status markers
    clause = re.sub(
        r"\b(DRAFT|CONFIDENTIAL|INTERNAL USE ONLY|UNCLASSIFIED|PROTECTED|"
        r"OFFICIAL|SECRET|TOP SECRET)\b",
        "",
        clause,
        flags=re.IGNORECASE,
    )

    # Collapse whitespace
    clause = re.sub(r"\s+", " ", clause)

    # Remove standalone numbers (stray page numbers)
    clause = re.sub(r"^\d+\.?\s*$", "", clause)

    # Normalise Unicode typography
    for src, dst in [
        ("\u2013", "-"), ("\u2014", "-"),
        ("\u2018", "'"), ("\u2019", "'"),
        ("\u201c", '"'), ("\u201d", '"'),
    ]:
        clause = clause.replace(src, dst)

    return clause.strip()
